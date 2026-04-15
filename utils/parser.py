"""
utils/parser.py
Extracts demographics and psychological test scores from PDF and DOCX files.
Handles messy, varied clinical report formats.
"""

import re
import io
from typing import Any

import pdfplumber
import fitz  # PyMuPDF
from docx import Document


# ─────────────────────────────────────────────
# DEMOGRAPHIC PATTERNS
# ─────────────────────────────────────────────

DEMO_PATTERNS = {
    "name": [
        r"(?:patient|client|name|الاسم|المريض)\s*[:\-]\s*([^\n\r,|]{2,50})",
        r"^Name\s*[:\-]\s*(.+)$",
    ],
    "age": [
        r"(?:age|عمر|العمر)\s*[:\-]\s*(\d{1,3})\s*(?:years?|سنة|سنوات)?",
        r"(\d{1,3})\s*(?:y\.?o\.?|years?\s*old)",
    ],
    "gender": [
        r"(?:gender|sex|الجنس|النوع)\s*[:\-]\s*(male|female|ذكر|أنثى|m|f)",
        r"\b(male|female|ذكر|أنثى)\b",
    ],
    "dob": [
        r"(?:dob|date\s*of\s*birth|تاريخ\s*الميلاد)\s*[:\-]\s*([\d/\-\.]{6,12})",
    ],
    "date_assessed": [
        r"(?:date\s*(?:of\s*)?(?:assessment|evaluation|testing)|تاريخ\s*التقييم)\s*[:\-]\s*([\d/\-\.]{6,12})",
        r"(?:assessed|evaluated)\s*(?:on)?\s*[:\-]?\s*([\d/\-\.]{6,12})",
    ],
    "referral_reason": [
        r"(?:referral\s*reason|reason\s*for\s*referral|سبب\s*الإحالة)\s*[:\-]\s*([^\n\r]{5,200})",
    ],
    "education": [
        r"(?:education|educational\s*level|المستوى\s*التعليمي|التعليم)\s*[:\-]\s*([^\n\r]{2,60})",
    ],
    "clinician": [
        r"(?:clinician|psychologist|examiner|الأخصائي|المعالج)\s*[:\-]\s*([^\n\r]{2,60})",
    ],
}

# Common psychological score keywords mapped to canonical test names
SCORE_KEYWORDS = {
    # IQ / Cognitive
    "full scale iq": "Full Scale IQ (FSIQ)",
    "fsiq": "Full Scale IQ (FSIQ)",
    "verbal comprehension": "Verbal Comprehension Index (VCI)",
    "vci": "Verbal Comprehension Index (VCI)",
    "perceptual reasoning": "Perceptual Reasoning Index (PRI)",
    "pri": "Perceptual Reasoning Index (PRI)",
    "working memory": "Working Memory Index (WMI)",
    "wmi": "Working Memory Index (WMI)",
    "processing speed": "Processing Speed Index (PSI)",
    "psi": "Processing Speed Index (PSI)",
    "fluid reasoning": "Fluid Reasoning Index (FRI)",
    "visual spatial": "Visual Spatial Index (VSI)",
    "general ability": "General Ability Index (GAI)",
    # Academic
    "reading": "Reading",
    "math": "Mathematics",
    "mathematics": "Mathematics",
    "written expression": "Written Expression",
    "spelling": "Spelling",
    # Memory
    "immediate memory": "Immediate Memory",
    "delayed memory": "Delayed Memory",
    "visual memory": "Visual Memory",
    "auditory memory": "Auditory Memory",
    # Adaptive / Behavioral
    "adaptive behavior": "Adaptive Behavior Composite",
    "conceptual": "Conceptual Domain",
    "social": "Social Domain",
    "practical": "Practical Domain",
    # Attention
    "inattention": "Inattention",
    "hyperactivity": "Hyperactivity/Impulsivity",
    "adhd": "ADHD Index",
    "attention": "Attention Index",
    # Emotional / Anxiety
    "anxiety": "Anxiety",
    "depression": "Depression",
    "externalizing": "Externalizing Problems",
    "internalizing": "Internalizing Problems",
    "total problems": "Total Problems",
    "thought problems": "Thought Problems",
    "social problems": "Social Problems",
    # Autism
    "social communication": "Social Communication",
    "restricted repetitive": "Restricted/Repetitive Behaviors",
    "ados": "ADOS-2 Total",
    # Executive Function
    "inhibit": "Inhibit",
    "shift": "Shift",
    "emotional control": "Emotional Control",
    "initiate": "Initiate",
    "working memory index": "Working Memory (BRIEF)",
    "plan": "Plan/Organize",
    "organization": "Organization of Materials",
    "monitor": "Monitor",
    "metacognition": "Metacognition Index",
    "behavioral regulation": "Behavioral Regulation Index",
    "global executive": "Global Executive Composite",
}


# ─────────────────────────────────────────────
# TEXT EXTRACTION
# ─────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Try pdfplumber first, fall back to PyMuPDF."""
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
            text = "\n".join(pages)
    except Exception:
        pass

    if len(text.strip()) < 100:
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            pages = [page.get_text() for page in doc]
            text = "\n".join(pages)
            doc.close()
        except Exception:
            pass

    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX including tables."""
    text_parts = []
    try:
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {e}")
    return "\n".join(text_parts)


def extract_raw_text(file_bytes: bytes, filename: str) -> str:
    fname = filename.lower()
    if fname.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif fname.endswith(".docx") or fname.endswith(".doc"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Unsupported file type. Please upload PDF or DOCX.")


# ─────────────────────────────────────────────
# DEMOGRAPHICS EXTRACTION
# ─────────────────────────────────────────────

def extract_demographics(text: str) -> dict:
    demographics = {}
    confidence = {}

    for field, patterns in DEMO_PATTERNS.items():
        found = None
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                found = match.group(1).strip()
                found = re.sub(r'\s+', ' ', found)
                break
        demographics[field] = found or ""
        confidence[field] = "high" if found else "low"

    return demographics, confidence


# ─────────────────────────────────────────────
# SCORE EXTRACTION
# ─────────────────────────────────────────────

def extract_scores(text: str) -> list[dict]:
    """
    Extract test scores robustly. Looks for patterns like:
      - 'Full Scale IQ: 95'
      - 'FSIQ  95  37th  Average'
      - table rows with score, percentile, classification
    Returns list of dicts with keys: test, score, percentile, classification, confidence
    """
    scores = []
    seen_tests = set()

    lines = text.split("\n")

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped or len(line_stripped) < 3:
            continue

        line_lower = line_stripped.lower()

        # Match known test keywords
        matched_test = None
        for keyword, canonical in SCORE_KEYWORDS.items():
            if keyword in line_lower:
                matched_test = canonical
                break

        if not matched_test:
            continue

        if matched_test in seen_tests:
            continue

        # Extract numeric score
        numbers = re.findall(r'\b(\d{1,3})\b', line_stripped)
        score_val = None
        percentile_val = None

        for num_str in numbers:
            num = int(num_str)
            # Scores are typically 40–160 for standard scores, 1–99 for percentiles
            if 40 <= num <= 160 and score_val is None:
                score_val = num
            elif 1 <= num <= 99 and score_val is not None and percentile_val is None:
                percentile_val = num

        # Extract classification
        classification = ""
        class_patterns = [
            r'\b(very superior|superior|high average|average|low average|'
            r'borderline|extremely low|below average|above average|'
            r'gifted|deficient|impaired|mild|moderate|severe|'
            r'ممتاز|متوسط|أعلى\s*من\s*المتوسط|أدنى\s*من\s*المتوسط|منخفض|مرتفع)\b'
        ]
        for cp in class_patterns:
            cm = re.search(cp, line_stripped, re.IGNORECASE)
            if cm:
                classification = cm.group(1).strip()
                break

        if score_val is None:
            # Try colon pattern: "Test Name: 95"
            colon_match = re.search(r':\s*(\d{2,3})', line_stripped)
            if colon_match:
                num = int(colon_match.group(1))
                if 40 <= num <= 160:
                    score_val = num

        confidence = "high" if score_val is not None else "low"

        scores.append({
            "test": matched_test,
            "score": score_val if score_val is not None else "",
            "percentile": percentile_val if percentile_val is not None else "",
            "classification": classification,
            "confidence": confidence,
        })
        seen_tests.add(matched_test)

    # Deduplicate by test name keeping best confidence
    final = {}
    for s in scores:
        t = s["test"]
        if t not in final or (s["confidence"] == "high" and final[t]["confidence"] == "low"):
            final[t] = s

    return list(final.values())


# ─────────────────────────────────────────────
# MAIN ENTRY POINT
# ─────────────────────────────────────────────

def extract_data_from_file(file_bytes: bytes, filename: str) -> dict:
    """
    Main extraction function.
    Returns:
      {
        "raw_text": str,
        "demographics": dict,
        "demo_confidence": dict,
        "scores": list[dict],
      }
    """
    raw_text = extract_raw_text(file_bytes, filename)

    if len(raw_text.strip()) < 50:
        raise ValueError(
            "Could not extract readable text from this file. "
            "Please ensure it is not a scanned image-only PDF."
        )

    demographics, demo_confidence = extract_demographics(raw_text)
    scores = extract_scores(raw_text)

    return {
        "raw_text": raw_text,
        "demographics": demographics,
        "demo_confidence": demo_confidence,
        "scores": scores,
    }
