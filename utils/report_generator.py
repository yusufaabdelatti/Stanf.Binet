"""
utils/report_generator.py
Generates professional PDF and DOCX clinical reports using ReportLab and python-docx.
"""

import io
import re
from datetime import datetime
from typing import Any

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    HRFlowable,
    Image as RLImage,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas as rl_canvas

# DOCX
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─────────────────────────────────────────────
# COLORS
# ─────────────────────────────────────────────

PRIMARY = colors.HexColor("#1e3a5f")
ACCENT = colors.HexColor("#2563eb")
LIGHT_BG = colors.HexColor("#f0f4ff")
BORDER = colors.HexColor("#e2e8f0")
TEXT = colors.HexColor("#1e293b")
MUTED = colors.HexColor("#64748b")
WHITE = colors.white
SUCCESS_BG = colors.HexColor("#f0fdf4")
WARNING_BG = colors.HexColor("#fffbeb")


# ─────────────────────────────────────────────
# SCORE CLASSIFICATION HELPERS
# ─────────────────────────────────────────────

def classify_score_color(score) -> colors.Color:
    """Return a color based on standard score range."""
    try:
        s = int(score)
    except (ValueError, TypeError):
        return MUTED
    if s >= 130:
        return colors.HexColor("#7c3aed")   # Very Superior – purple
    elif s >= 120:
        return colors.HexColor("#2563eb")   # Superior – blue
    elif s >= 110:
        return colors.HexColor("#0891b2")   # High Average – teal
    elif s >= 90:
        return colors.HexColor("#16a34a")   # Average – green
    elif s >= 80:
        return colors.HexColor("#d97706")   # Low Average – amber
    elif s >= 70:
        return colors.HexColor("#ea580c")   # Borderline – orange
    else:
        return colors.HexColor("#dc2626")   # Extremely Low – red


# ─────────────────────────────────────────────
# PAGE NUMBER CANVAS
# ─────────────────────────────────────────────

class NumberedCanvas(rl_canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            rl_canvas.Canvas.showPage(self)
        rl_canvas.Canvas.save(self)

    def draw_page_number(self, page_count):
        self.setFont("Helvetica", 8)
        self.setFillColor(MUTED)
        self.drawRightString(
            A4[0] - 2 * cm,
            1.2 * cm,
            f"Page {self._pageNumber} of {page_count}",
        )
        self.setStrokeColor(BORDER)
        self.setLineWidth(0.5)
        self.line(2 * cm, 1.6 * cm, A4[0] - 2 * cm, 1.6 * cm)


# ─────────────────────────────────────────────
# STYLES
# ─────────────────────────────────────────────

def get_styles(is_rtl: bool = False) -> dict:
    alignment = TA_RIGHT if is_rtl else TA_LEFT

    styles = {
        "title": ParagraphStyle(
            "ReportTitle",
            fontName="Helvetica-Bold",
            fontSize=22,
            textColor=PRIMARY,
            spaceAfter=4,
            alignment=TA_CENTER,
        ),
        "subtitle": ParagraphStyle(
            "Subtitle",
            fontName="Helvetica",
            fontSize=11,
            textColor=MUTED,
            spaceAfter=2,
            alignment=TA_CENTER,
        ),
        "section_header": ParagraphStyle(
            "SectionHeader",
            fontName="Helvetica-Bold",
            fontSize=13,
            textColor=PRIMARY,
            spaceBefore=14,
            spaceAfter=6,
            alignment=alignment,
        ),
        "body": ParagraphStyle(
            "Body",
            fontName="Helvetica",
            fontSize=10,
            textColor=TEXT,
            leading=16,
            spaceAfter=6,
            alignment=TA_JUSTIFY if not is_rtl else TA_RIGHT,
        ),
        "label": ParagraphStyle(
            "Label",
            fontName="Helvetica-Bold",
            fontSize=9,
            textColor=MUTED,
            spaceAfter=2,
            alignment=alignment,
        ),
        "value": ParagraphStyle(
            "Value",
            fontName="Helvetica",
            fontSize=10,
            textColor=TEXT,
            spaceAfter=2,
            alignment=alignment,
        ),
        "footer": ParagraphStyle(
            "Footer",
            fontName="Helvetica",
            fontSize=8,
            textColor=MUTED,
            alignment=TA_CENTER,
        ),
        "disclaimer": ParagraphStyle(
            "Disclaimer",
            fontName="Helvetica-Oblique",
            fontSize=8,
            textColor=MUTED,
            leading=12,
            alignment=TA_CENTER,
        ),
    }
    return styles


# ─────────────────────────────────────────────
# BUILD DEMOGRAPHICS TABLE
# ─────────────────────────────────────────────

DEMO_FIELD_LABELS_EN = {
    "name": "Patient Name",
    "age": "Age",
    "gender": "Gender",
    "dob": "Date of Birth",
    "date_assessed": "Date Assessed",
    "referral_reason": "Referral Reason",
    "education": "Education Level",
    "clinician": "Clinician",
}

DEMO_FIELD_LABELS_AR = {
    "name": "اسم المريض",
    "age": "العمر",
    "gender": "الجنس",
    "dob": "تاريخ الميلاد",
    "date_assessed": "تاريخ التقييم",
    "referral_reason": "سبب الإحالة",
    "education": "المستوى التعليمي",
    "clinician": "الأخصائي",
}


def build_demographics_table(demographics: dict, language: str, styles: dict):
    labels = DEMO_FIELD_LABELS_AR if language == "Arabic" else DEMO_FIELD_LABELS_EN
    order = ["name", "age", "gender", "dob", "date_assessed", "education", "referral_reason", "clinician"]

    rows = []
    for field in order:
        val = demographics.get(field, "")
        if not val or str(val).strip() == "":
            val = "—"
        label_text = labels.get(field, field)
        rows.append([
            Paragraph(f"<b>{label_text}</b>", styles["label"]),
            Paragraph(str(val), styles["value"]),
        ])

    # Split into two columns of rows
    half = (len(rows) + 1) // 2
    left_rows = rows[:half]
    right_rows = rows[half:]

    # Pad
    while len(right_rows) < len(left_rows):
        right_rows.append([Paragraph("", styles["label"]), Paragraph("", styles["value"])])

    combined = []
    for l, r in zip(left_rows, right_rows):
        combined.append(l + r)

    col_widths = [4.5 * cm, 6.5 * cm, 4.5 * cm, 6.5 * cm]
    t = Table(combined, colWidths=col_widths)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), LIGHT_BG),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [WHITE, LIGHT_BG]),
        ("GRID", (0, 0), (-1, -1), 0.5, BORDER),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("ROUNDEDCORNERS", [6, 6, 6, 6]),
    ]))
    return t


# ─────────────────────────────────────────────
# BUILD SCORES TABLE
# ─────────────────────────────────────────────

def build_scores_table(scores: list[dict], language: str, styles: dict):
    if language == "Arabic":
        headers = ["الاختبار / المقياس", "الدرجة", "الشريحة المئوية", "التصنيف"]
    else:
        headers = ["Test / Scale", "Score", "Percentile", "Classification"]

    header_row = [Paragraph(f"<b>{h}</b>", styles["label"]) for h in headers]

    data = [header_row]
    for s in scores:
        score_val = s.get("score", "")
        score_color = classify_score_color(score_val)
        score_text = str(score_val) if score_val != "" else "—"
        pct_text = str(s.get("percentile", "")) if s.get("percentile", "") != "" else "—"
        cls_text = s.get("classification", "") or "—"

        data.append([
            Paragraph(str(s.get("test", "")), styles["body"]),
            Paragraph(f"<b><font color='#{score_color.hexval()[2:]}'>{score_text}</font></b>", styles["body"]),
            Paragraph(pct_text, styles["body"]),
            Paragraph(cls_text, styles["body"]),
        ])

    col_widths = [9 * cm, 2.5 * cm, 3.5 * cm, 7 * cm]
    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), PRIMARY),
        ("TEXTCOLOR", (0, 0), (-1, 0), WHITE),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, LIGHT_BG]),
        ("GRID", (0, 0), (-1, -1), 0.5, BORDER),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("ALIGN", (1, 0), (2, -1), "CENTER"),
    ]))
    return t


# ─────────────────────────────────────────────
# EMBED CHART IMAGE
# ─────────────────────────────────────────────

def embed_chart(fig, width_cm: float = 17, height_cm: float = 9):
    """Convert a Plotly figure to a ReportLab Image."""
    try:
        img_bytes = fig.to_image(format="png", width=900, height=480, scale=2)
        return RLImage(io.BytesIO(img_bytes), width=width_cm * cm, height=height_cm * cm)
    except Exception:
        return None


# ─────────────────────────────────────────────
# INTERPRETATION FORMATTER
# ─────────────────────────────────────────────

def format_interpretation(text: str, styles: dict) -> list:
    """Split interpretation into styled paragraphs for the PDF."""
    flowables = []
    lines = text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            flowables.append(Spacer(1, 4))
            continue
        # Detect section headers (numbered or ALL CAPS or ends with colon)
        if re.match(r'^(#{1,3}|\d+[\.\)]|\*{2})', stripped) or (
            stripped.isupper() and len(stripped) > 3
        ) or stripped.endswith(":"):
            clean = re.sub(r'^[#\d\.\)\*\s]+', '', stripped).rstrip(":")
            flowables.append(Spacer(1, 6))
            flowables.append(Paragraph(clean, styles["section_header"]))
        else:
            # Remove markdown bold/italic
            clean = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', stripped)
            flowables.append(Paragraph(clean, styles["body"]))
    return flowables


# ─────────────────────────────────────────────
# GENERATE PDF
# ─────────────────────────────────────────────

def generate_pdf_report(
    demographics: dict,
    scores: list[dict],
    interpretation: str,
    language: str,
    center_name: str,
    clinician_name: str,
    logo_bytes: bytes | None,
    bar_fig=None,
    radar_fig=None,
) -> bytes:
    """
    Generate a full professional PDF report.
    Returns PDF as bytes.
    """
    buf = io.BytesIO()
    is_rtl = language == "Arabic"
    styles = get_styles(is_rtl)

    # Document margins
    doc = BaseDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        title=f"Psychological Assessment Report — {demographics.get('name', '')}",
        author=clinician_name or center_name or "PsychReport AI",
    )

    frame = Frame(
        doc.leftMargin,
        doc.bottomMargin,
        doc.width,
        doc.height,
        id="normal",
    )
    template = PageTemplate(id="main", frames=frame)
    doc.addPageTemplates([template])

    story = []

    # ── HEADER ──────────────────────────────
    # Logo
    if logo_bytes:
        try:
            logo_img = RLImage(io.BytesIO(logo_bytes), width=3.5 * cm, height=3.5 * cm)
            logo_img.hAlign = "CENTER"
            story.append(logo_img)
            story.append(Spacer(1, 8))
        except Exception:
            pass

    # Center name
    if center_name:
        story.append(Paragraph(center_name, styles["title"]))
        story.append(Spacer(1, 4))

    # Report title
    title_text = "تقرير التقييم النفسي" if language == "Arabic" else "Psychological Assessment Report"
    story.append(Paragraph(title_text, styles["title"] if not center_name else styles["subtitle"]))
    story.append(Spacer(1, 4))

    # Date generated
    gen_date = datetime.now().strftime("%B %d, %Y")
    date_label = "تاريخ إنشاء التقرير" if language == "Arabic" else "Report Generated"
    story.append(Paragraph(f"{date_label}: {gen_date}", styles["subtitle"]))

    if clinician_name:
        clinician_label = "الأخصائي" if language == "Arabic" else "Prepared by"
        story.append(Paragraph(f"{clinician_label}: {clinician_name}", styles["subtitle"]))

    story.append(Spacer(1, 10))
    story.append(HRFlowable(width="100%", thickness=2, color=ACCENT, spaceAfter=12))

    # ── DEMOGRAPHICS ────────────────────────
    demo_label = "البيانات الديموغرافية" if language == "Arabic" else "Patient Information"
    story.append(Paragraph(demo_label, styles["section_header"]))
    story.append(Spacer(1, 4))
    story.append(build_demographics_table(demographics, language, styles))
    story.append(Spacer(1, 16))

    # ── SCORES TABLE ────────────────────────
    if scores:
        scores_label = "نتائج الاختبارات" if language == "Arabic" else "Test Results Summary"
        story.append(Paragraph(scores_label, styles["section_header"]))
        story.append(Spacer(1, 4))
        story.append(build_scores_table(scores, language, styles))
        story.append(Spacer(1, 16))

    # ── CHARTS ──────────────────────────────
    if bar_fig is not None:
        chart_label = "التمثيل البياني" if language == "Arabic" else "Score Visualization"
        story.append(Paragraph(chart_label, styles["section_header"]))
        story.append(Spacer(1, 6))
        chart_img = embed_chart(bar_fig)
        if chart_img:
            story.append(chart_img)
            story.append(Spacer(1, 8))

    if radar_fig is not None:
        radar_img = embed_chart(radar_fig)
        if radar_img:
            story.append(radar_img)
            story.append(Spacer(1, 16))

    # ── INTERPRETATION ──────────────────────
    if interpretation:
        interp_label = "التفسير السريري" if language == "Arabic" else "Clinical Interpretation"
        story.append(PageBreak())
        story.append(Paragraph(interp_label, styles["section_header"]))
        story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER, spaceAfter=8))
        story.extend(format_interpretation(interpretation, styles))
        story.append(Spacer(1, 20))

    # ── DISCLAIMER ──────────────────────────
    story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER, spaceBefore=12, spaceAfter=8))
    disclaimer_en = (
        "This report is intended for professional use only. "
        "The findings presented herein are based solely on the scores provided and "
        "should be interpreted by a qualified clinician in conjunction with all available clinical information."
    )
    disclaimer_ar = (
        "هذا التقرير مخصص للاستخدام المهني فقط. "
        "تستند النتائج الواردة فيه إلى الدرجات المُقدَّمة حصرًا، "
        "وينبغي تفسيرها من قِبَل متخصص مؤهل في ضوء جميع المعلومات السريرية المتاحة."
    )
    story.append(Paragraph(disclaimer_ar if language == "Arabic" else disclaimer_en, styles["disclaimer"]))

    # Build PDF
    doc.build(story, canvasmaker=NumberedCanvas)
    return buf.getvalue()


# ─────────────────────────────────────────────
# GENERATE DOCX
# ─────────────────────────────────────────────

def generate_docx_report(
    demographics: dict,
    scores: list[dict],
    interpretation: str,
    language: str,
    center_name: str,
    clinician_name: str,
    logo_bytes: bytes | None,
) -> bytes:
    """Generate a professional DOCX report."""
    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    is_rtl = language == "Arabic"
    align_val = WD_ALIGN_PARAGRAPH.RIGHT if is_rtl else WD_ALIGN_PARAGRAPH.LEFT

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        h.alignment = align_val
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    def add_para(text, bold=False, italic=False, size=10):
        p = doc.add_paragraph()
        p.alignment = align_val
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        return p

    # Logo
    if logo_bytes:
        try:
            doc.add_picture(io.BytesIO(logo_bytes), width=Inches(1.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    # Title
    if center_name:
        title_p = doc.add_heading(center_name, level=1)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    report_title = "تقرير التقييم النفسي" if language == "Arabic" else "Psychological Assessment Report"
    h = doc.add_heading(report_title, level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    gen_date = datetime.now().strftime("%B %d, %Y")
    date_label = "تاريخ إنشاء التقرير" if language == "Arabic" else "Report Generated"
    add_para(f"{date_label}: {gen_date}", size=10).alignment = WD_ALIGN_PARAGRAPH.CENTER

    if clinician_name:
        clinician_label = "الأخصائي" if language == "Arabic" else "Prepared by"
        add_para(f"{clinician_label}: {clinician_name}", bold=True, size=10).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Demographics
    demo_label = "البيانات الديموغرافية" if language == "Arabic" else "Patient Information"
    add_heading(demo_label, level=2)
    labels = DEMO_FIELD_LABELS_AR if language == "Arabic" else DEMO_FIELD_LABELS_EN
    order = ["name", "age", "gender", "dob", "date_assessed", "education", "referral_reason", "clinician"]

    demo_table = doc.add_table(rows=1, cols=2)
    demo_table.style = "Table Grid"
    hdr = demo_table.rows[0].cells
    hdr[0].text = "Field" if language == "English" else "الحقل"
    hdr[1].text = "Value" if language == "English" else "القيمة"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell._tc.get_or_add_tcPr()

    for field in order:
        val = demographics.get(field, "") or "—"
        row = demo_table.add_row().cells
        row[0].text = labels.get(field, field)
        row[1].text = str(val)

    doc.add_paragraph()

    # Scores table
    if scores:
        scores_label = "نتائج الاختبارات" if language == "Arabic" else "Test Results Summary"
        add_heading(scores_label, level=2)

        score_tbl = doc.add_table(rows=1, cols=4)
        score_tbl.style = "Table Grid"
        if language == "Arabic":
            headers = ["الاختبار", "الدرجة", "الشريحة المئوية", "التصنيف"]
        else:
            headers = ["Test / Scale", "Score", "Percentile", "Classification"]
        hdr_cells = score_tbl.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for para in hdr_cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True

        for s in scores:
            row = score_tbl.add_row().cells
            row[0].text = str(s.get("test", ""))
            row[1].text = str(s.get("score", "")) or "—"
            row[2].text = str(s.get("percentile", "")) if s.get("percentile", "") != "" else "—"
            row[3].text = str(s.get("classification", "")) or "—"

        doc.add_paragraph()

    # Interpretation
    if interpretation:
        interp_label = "التفسير السريري" if language == "Arabic" else "Clinical Interpretation"
        add_heading(interp_label, level=2)
        for line in interpretation.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue
            clean = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', stripped)
            p = doc.add_paragraph(clean)
            p.alignment = align_val

    # Disclaimer
    doc.add_paragraph()
    disclaimer = (
        "This report is for professional use only and should be interpreted by a qualified clinician."
        if language == "English"
        else "هذا التقرير للاستخدام المهني فقط وينبغي تفسيره من قِبَل متخصص مؤهل."
    )
    add_para(disclaimer, italic=True, size=9).alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
